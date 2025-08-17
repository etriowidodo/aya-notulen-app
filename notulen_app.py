import asyncio
import logging
import sys
import tempfile
import threading
import queue
import time
import datetime
import os
from dataclasses import dataclass, field
from typing import List, Dict, Optional
import numpy as np
import sounddevice as sd
import webrtcvad
import torch
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from torch.nn import functional as F
import torchaudio
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from faster_whisper import WhisperModel
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog, filedialog
import sqlite3
from speechbrain.inference import EncoderClassifier
from collections import Counter
from transformers import pipeline, T5Config
import os
os.environ["PYTORCH_CUDA_ALLOC_CONF"] = "expandable_segments:True"
# Configure logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# Constants
SAMPLE_RATE = 16000
FRAME_MS = 30
VAD_MODE = 2
MIN_UTTER_SEC = 0.9
MAX_UTTER_SEC = 20.0
SILENCE_TAIL_FRAMES = 8
SPEAKER_SIM_THRESHOLD = 0.80
WHISPER_MODEL_SIZE = "medium"
WHISPER_DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
WHISPER_COMPUTE_TYPE = "float16" if WHISPER_DEVICE == "cuda" else "int8"
DB_FILE = "notulen.db"

# Data Structures
@dataclass
class Utterance:
    speaker_id: str
    text: str
    start_time: float
    end_time: float

@dataclass
class MeetingState:
    running: bool = False
    start_time: Optional[datetime.datetime] = None
    end_time: Optional[datetime.datetime] = None
    utterances: List[Utterance] = field(default_factory=list)
    speaker_map: Dict[str, str] = field(default_factory=dict)
    speaker_embeds: Dict[str, torch.Tensor] = field(default_factory=dict)
    next_speaker_idx: int = 1

# Audio Device Management
class AudioDeviceManager:
    @staticmethod
    def get_input_devices():
        devices = []
        for idx, dev in enumerate(sd.query_devices()):
            if dev['max_input_channels'] > 0:
                devices.append((idx, dev['name']))
        return devices

    @staticmethod
    def get_loopback_devices():
        devices = []
        default_output = sd.default.device[1]
        output_name = sd.query_devices(default_output)['name']
        for idx, dev in enumerate(sd.query_devices()):
            if "(loopback)" in dev['name'].lower() and output_name.split('(')[0].strip() in dev['name']:
                devices.append((idx, dev['name']))
        return devices

    @staticmethod
    def get_default_input():
        default_idx = sd.default.device[0]
        return default_idx, sd.query_devices(default_idx)['name']

# Mock helper functions (since not provided)
def find_speaker_by_embedding(emb, threshold=0.8):
    return None

def save_speaker_embedding(name, emb):
    pass

def init_speaker_db():
    pass

def merge_speakers(main, sec):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("UPDATE transkrip SET pembicara=? WHERE pembicara=?", (main, sec))
    conn.commit()
    conn.close()

# Meeting Analyzer Classes
class MeetingAnalyzer:
    def __init__(self):
        try:
            model_path = self._get_valid_model_path(
                't5-indonesian-summarization',
                ['config.json', 'model.safetensors', 'tokenizer_config.json']
            )
            config = T5Config.from_pretrained(model_path, forced_bos_token_id=0)
            self.summarizer = pipeline(
                "summarization",
                model=model_path,
                config=config,
                device=WHISPER_DEVICE
            )
            logging.info("Loaded local MeetingAnalyzer model")
        except Exception as e:
            logging.error(f"Error loading local model: {e}")
            self.summarizer = lambda x, **kwargs: [{"summary_text": "Mock summary"}]

    def _get_valid_model_path(self, model_name, required_files):
        possible_paths = [
            os.path.join(os.path.dirname(__file__), 'models', model_name),
            os.path.join(getattr(sys, '_MEIPASS', ''), 'models', model_name),
            os.path.join('models', model_name),
            os.path.join('.', 'models', model_name)
        ]
        for path in possible_paths:
            if path and os.path.exists(path):
                if all(os.path.exists(os.path.join(path, f)) for f in required_files):
                    return path
        raise FileNotFoundError(f"Model {model_name} not found in {possible_paths}")

    def _chunk_text(self, text, max_chars=1024):
        chunks = []
        while len(text) > max_chars:
            cut = text[:max_chars]
            last_dot = cut.rfind(".")
            if last_dot == -1:
                last_dot = max_chars
            chunks.append(text[:last_dot+1].strip())
            text = text[last_dot+1:].strip()
        if text:
            chunks.append(text)
        return chunks

    def generate_summary(self, meeting):
        if not meeting.utterances:
            return "Tidak ada transkrip yang dapat diringkas"
        full_text = " ".join([utt.text for utt in meeting.utterances])
        chunks = self._chunk_text(full_text)
        partial_summaries = []
        for ch in chunks:
            result = self.summarizer(ch, max_length=130, min_length=50, do_sample=False)
            partial_summaries.append(result[0]['summary_text'])
        combined_summary = " ".join(partial_summaries)
        final_points = self.summarizer(
            "Ringkas poin-poin utama dari teks berikut:\n" + combined_summary,
            max_length=150,
            min_length=60,
            do_sample=False
        )[0]['summary_text']
        return "=== RINGKASAN RAPAT ===\n" + final_points

    def generate_conclusion(self, summary):
        conclusion = self.summarizer(
            "Buat kesimpulan singkat dari ringkasan rapat berikut:\n" + summary,
            max_length=100,
            min_length=40,
            do_sample=False
        )[0]['summary_text']
        return "=== KESIMPULAN ===\n" + conclusion

class FullMeetingAnalyzer:
    def __init__(self):
        try:
            model_path = self._get_valid_model_path(
                't5-indonesian-summarization',
                ['config.json', 'model.safetensors', 'tokenizer_config.json']
            )
            config = T5Config.from_pretrained(model_path, forced_bos_token_id=0)
            self.summarizer = pipeline(
                "summarization",
                model=model_path,
                config=config,
                device=WHISPER_DEVICE
            )
            logging.info("Loaded local FullMeetingAnalyzer model")
        except Exception as e:
            logging.error(f"Error loading local model: {e}")
            self.summarizer = lambda x, **kwargs: [{"summary_text": "Mock full summary"}]

    def _get_valid_model_path(self, model_name, required_files):
        possible_paths = [
            os.path.join(os.path.dirname(__file__), 'models', model_name),
            os.path.join(getattr(sys, '_MEIPASS', ''), 'models', model_name),
            os.path.join('models', model_name),
            os.path.join('.', 'models', model_name)
        ]
        for path in possible_paths:
            if path and os.path.exists(path):
                if all(os.path.exists(os.path.join(path, f)) for f in required_files):
                    return path
        raise FileNotFoundError(f"Model {model_name} not found in {possible_paths}")

    def _merge_to_paragraphs(self, meeting):
        paragraphs = []
        current_para = []
        for utt in meeting.utterances:
            text = utt.text.strip()
            if len(text.split()) < 4:
                continue
            current_para.append(text)
        if current_para:
            paragraphs.append(" ".join(current_para))
        return paragraphs

    def generate_summary(self, meeting):
        if not meeting.utterances:
            return "Tidak ada transkrip yang dapat diringkas"
        paragraphs = self._merge_to_paragraphs(meeting)
        if not paragraphs:
            return "Tidak ada materi penting untuk diringkas"
        full_text = " ".join(paragraphs)
        if self.summarizer:
            try:
                result = self.summarizer(
                    full_text,
                    max_length=130,
                    min_length=50,
                    do_sample=False,
                    num_beams=4,
                    length_penalty=1.2,
                    no_repeat_ngram_size=3
                )
                summary_text = result[0]['summary_text']
            except Exception as e:
                logging.error(f"Summarizer error: {e}")
                summary_text = full_text
        else:
            summary_text = full_text
        lines = [
            "=== RINGKASAN RAPAT ===",
            f"Tanggal: {meeting.start_time.strftime('%d/%m/%Y')}",
            f"Waktu: {meeting.start_time.strftime('%H:%M')} - {meeting.end_time.strftime('%H:%M')}",
            f"Durasi: {(meeting.end_time - meeting.start_time).total_seconds() / 60:.1f} menit",
            "",
            summary_text.strip(),
            "",
            "=== PARTISIPASI PEMBICARA ==="
        ]
        speaker_counts = Counter()
        for utt in meeting.utterances:
            spk = meeting.speaker_map.get(utt.speaker_id, utt.speaker_id)
            speaker_counts[spk] += 1
        for spk, cnt in speaker_counts.most_common():
            lines.append(f"- {spk}: {cnt} interaksi")
        return "\n".join(lines)

# Meeting Exporter with Improved Error Handling
class MeetingExporter:
    def __init__(self, engine, ent_judul, ent_peserta, render_full_transcript_text, progress_callback=None):
        self.engine = engine
        self.ent_judul = ent_judul
        self.ent_peserta = ent_peserta
        self.render_full_transcript_text = render_full_transcript_text
        self.progress_callback = progress_callback
        self.lock = threading.Lock()  # Add a lock for thread safety

    async def update_progress(self, value, max_value):
        """Thread-safe progress update"""
        if self.progress_callback:
            percentage = min(100, int((value / max_value) * 100))
            if hasattr(self.progress_callback, '__call__'):
                self.progress_callback(percentage)
            await asyncio.sleep(0)

    async def export_pdf(self):
        with self.lock:
            if not self.engine or not self.engine.meeting.utterances:
                await self.show_warning("Belum ada data rapat")
                return
        logging.info("Starting PDF export")
        torch.cuda.empty_cache()
        try:
            meeting = self.engine.meeting
            analyzer = MeetingAnalyzer()
            fullAnalyzer = FullMeetingAnalyzer()
            await self.update_progress(10, 100)
            judul = self.ent_judul.get().strip() or "Notulen Rapat"
            peserta = self.ent_peserta.get().strip()
            # Simpan transkrip ke file sementara
            with tempfile.NamedTemporaryFile(mode='w', delete=False, encoding='utf-8') as temp_file:
                for utt in meeting.utterances:
                    name = meeting.speaker_map.get(utt.speaker_id, utt.speaker_id)
                    temp_file.write(f"{name}: {utt.text}\n")
                temp_file_path = temp_file.name
            summary = analyzer.generate_summary(meeting)
            summary = analyzer.generate_conclusion(summary)
            fullsummary = fullAnalyzer.generate_summary(meeting)
            summary = summary + "\n\n=== RINGKASAN NARATIF ===\n" + fullsummary
            filename = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF Files", "*.pdf")],
                initialfile=f"Notulen_{meeting.start_time.strftime('%Y%m%d')}.pdf"
            )
            if not filename:
                logging.info("PDF export canceled by user")
                os.unlink(temp_file_path)
                return
            try:
                logging.debug(f"Saving PDF to {filename}")
                await self.update_progress(40, 100)
                doc = SimpleDocTemplate(
                    filename,
                    pagesize=A4,
                    rightMargin=2 * cm,
                    leftMargin=2 * cm,
                    topMargin=2 * cm,
                    bottomMargin=2 * cm
                )
                styles = getSampleStyleSheet()
                styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY))
                story = []
                await self.update_progress(50, 100)
                story.append(Paragraph(judul, styles['Title']))
                story.append(Spacer(1, 12))
                story.append(Paragraph(f"Tanggal: {meeting.start_time.strftime('%d %B %Y')}", styles['Normal']))
                story.append(Paragraph(
                    f"Waktu: {meeting.start_time.strftime('%H:%M')} - {meeting.end_time.strftime('%H:%M')}",
                    styles['Normal']
                ))
                if peserta:
                    story.append(Paragraph(f"Peserta: {peserta}", styles['Normal']))
                story.append(Spacer(1, 12))
                await self.update_progress(60, 100)
                story.append(Paragraph("<b>Ringkasan Rapat</b>", styles['Heading2']))
                for line in summary.split("\n"):
                    if line.strip():
                        story.append(Paragraph(line.strip(), styles['Justify']))
                story.append(PageBreak())
                await self.update_progress(80, 100)
                story.append(Paragraph("<b>Transkrip Lengkap</b>", styles['Heading2']))
                with open(temp_file_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        if ':' in line:
                            speaker, content = line.split(':', 1)
                            story.append(Paragraph(f"<b>{speaker}:</b> {content.strip()}", styles['Justify']))
                        else:
                            story.append(Paragraph(line.strip(), styles['Justify']))
                await self.update_progress(90, 100)
                story.append(PageBreak())
                story.append(Paragraph(
                    f"Dokumen ini dihasilkan otomatis pada {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}",
                    styles['Normal']
                ))
                story.append(Paragraph("Sistem Notulen AYA - Aplikasi Notulen Anak Bangsa", styles['Normal']))
                await self.update_progress(100, 100)
                doc.build(story)
                logging.info(f"PDF successfully saved to {filename}")
                messagebox.showinfo("Sukses", f"PDF berhasil disimpan:\n{filename}")
            except Exception as e:
                await self.show_error(f"Gagal menyimpan PDF: {str(e)}")
            finally:
                os.unlink(temp_file_path)  # Hapus file sementara
        except Exception as e:
            await self.show_error(f"Gagal menyimpan PDF: {str(e)}")
        finally:
            del analyzer
            del fullAnalyzer
            torch.cuda.empty_cache()
            import gc
            gc.collect()

    async def export_word(self):
        logging.info("Starting Word export")
        torch.cuda.empty_cache()
        if not self.engine or not self.engine.meeting.utterances:
            await self.show_warning("Belum ada data rapat")
            return
        try:
            meeting = self.engine.meeting
            analyzer = MeetingAnalyzer()
            fullAnalyzer = FullMeetingAnalyzer()
            await self.update_progress(10, 100)
            judul = self.ent_judul.get().strip() or "Notulen Rapat"
            peserta = self.ent_peserta.get().strip()
            # Simpan transkrip ke file sementara
            with tempfile.NamedTemporaryFile(mode='w', delete=False, encoding='utf-8') as temp_file:
                for utt in meeting.utterances:
                    name = meeting.speaker_map.get(utt.speaker_id, utt.speaker_id)
                    temp_file.write(f"{name}: {utt.text}\n")
                temp_file_path = temp_file.name
            summary_points = analyzer.generate_summary(meeting)
            summary_points = analyzer.generate_conclusion(summary_points)
            summary_narr = fullAnalyzer.generate_summary(meeting)
            await self.update_progress(30, 100)
            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                initialfile=f"Notulen_{meeting.start_time.strftime('%Y%m%d')}.docx"
            )
            if not filename:
                logging.info("Word export canceled by user")
                os.unlink(temp_file_path)
                return
            try:
                logging.debug(f"Saving Word to {filename}")
                await self.update_progress(40, 100)
                doc = Document()
                await self.update_progress(50, 100)
                styles = doc.styles
                if 'Body Small' not in styles:
                    s = styles.add_style('Body Small', WD_STYLE_TYPE.PARAGRAPH)
                    s.font.size = Pt(10)
                    s.font.name = 'Calibri'
                h = doc.add_heading(judul, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                meta = doc.add_paragraph()
                meta_run = meta.add_run(
                    f"Tanggal: {meeting.start_time.strftime('%d %B %Y')}\n"
                    f"Waktu: {meeting.start_time.strftime('%H:%M')} - {meeting.end_time.strftime('%H:%M')}\n"
                    f"Durasi: {(meeting.end_time - meeting.start_time).total_seconds() / 60:.1f} menit"
                )
                meta.style = doc.styles['Body Small']
                if peserta:
                    p = doc.add_paragraph()
                    p.add_run(f"Peserta: {peserta}")
                doc.add_paragraph()
                await self.update_progress(60, 100)
                doc.add_heading("Ringkasan (Poin-Poin Utama)", level=2)
                poin_lines = []
                for line in summary_points.splitlines():
                    s = line.strip()
                    if not s or s.startswith("===") or s.lower().startswith(("tanggal:", "waktu:", "durasi:")):
                        continue
                    poin_lines.append(s)
                for s in poin_lines:
                    if s[0:2].isdigit() or s[:2].replace('.', '').isdigit():
                        doc.add_paragraph(s, style='List Number')
                    else:
                        doc.add_paragraph(s, style='List Bullet')
                doc.add_paragraph()
                await self.update_progress(70, 100)
                doc.add_heading("Ringkasan Naratif", level=2)
                for line in summary_narr.splitlines():
                    s = line.strip()
                    if not s or s.startswith("===") or s.lower().startswith(("tanggal:", "waktu:", "durasi:")):
                        continue
                    doc.add_paragraph(s)
                doc.add_paragraph()
                await self.update_progress(80, 100)
                if "=== PARTISIPASI PEMBICARA ===" in summary_narr:
                    doc.add_heading("Partisipasi Pembicara", level=2)
                    part = False
                    for line in summary_narr.splitlines():
                        if "=== PARTISIPASI PEMBICARA ===" in line:
                            part = True
                            continue
                        if part:
                            s = line.strip()
                            if not s:
                                continue
                            doc.add_paragraph(s, style='List Bullet')
                    doc.add_paragraph()
                await self.update_progress(90, 100)
                doc.add_heading("Transkrip Lengkap", level=2)
                with open(temp_file_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        if ':' in line:
                            speaker, content = line.split(':', 1)
                            para = doc.add_paragraph()
                            run1 = para.add_run(f"{speaker.strip()}: ")
                            run1.bold = True
                            para.add_run(content.strip())
                        else:
                            doc.add_paragraph(line.strip())
                await self.update_progress(95, 100)
                doc.add_page_break()
                f = doc.add_paragraph()
                f.add_run(
                    f"Dokumen ini dihasilkan otomatis pada {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
                    f"Sistem Notulen AYA - Aplikasi Notulen Anak Bangsa"
                ).italic = True
                await self.update_progress(100, 100)
                doc.save(filename)
                logging.info(f"Word successfully saved to {filename}")
                messagebox.showinfo("Sukses", f"Word berhasil disimpan:\n{filename}")
            except Exception as e:
                await self.show_error(f"Gagal menyimpan Word: {str(e)}")
            finally:
                os.unlink(temp_file_path)  # Hapus file sementara
        except Exception as e:
            await self.show_error(f"Gagal menyimpan Word: {str(e)}")
        finally:
            del analyzer
            del fullAnalyzer
            torch.cuda.empty_cache()
            import gc
            gc.collect()
    async def show_warning(self, message):
        """Thread-safe warning message"""

        def show():
            messagebox.showwarning("Info", message)

        self.engine.gui.after(0, show)
        await asyncio.sleep(0)
    async def show_error(self, message):
        """Thread-safe error message"""

        def show():
            messagebox.showerror("Error", message)

        self.engine.gui.after(0, show)
        await asyncio.sleep(0)
# Splash Screen
class SplashScreen:
    def __init__(self, root, on_complete):
        self.root = root
        self.on_complete = on_complete
        self.window = tk.Toplevel(root)
        self.window.title("Loading AYA Notulen Rapat")
        self.window.geometry("400x200")
        self.window.transient(root)
        self.window.grab_set()
        self.window.resizable(False, False)

        # Center the splash screen
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f"{width}x{height}+{x}+{y}")

        # Try to set icon
        try:
            icon_path = self._get_resource_path('AYA.ico')
            self.window.iconbitmap(icon_path)
        except Exception as e:
            logging.warning(f"Failed to load splash icon: {str(e)}")

        # Splash screen content
        frame = ttk.Frame(self.window)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        ttk.Label(frame, text="AYA Notulen Rapat", font=("Helvetica", 16, "bold")).pack(pady=10)
        ttk.Label(frame, text="Memuat model dan inisialisasi...").pack(pady=5)
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X, padx=10, pady=10)

        # Start loading in a separate thread
        self.load_complete = False
        threading.Thread(target=self.load_resources, daemon=True).start()
        self.update_progress()

    def _get_resource_path(self, filename):
        base_path = os.path.dirname(os.path.abspath(__file__))
        possible_paths = [
            os.path.join(base_path, filename),
            os.path.join(getattr(sys, '_MEIPASS', ''), filename),
            filename
        ]
        for path in possible_paths:
            if path and os.path.exists(path):
                return path
        return None

    def load_resources(self):
        steps = [
            ("Inisialisasi database", 20),
            ("Memuat model MeetingAnalyzer", 60),
            ("Memuat model FullMeetingAnalyzer", 100)
        ]
        try:
            for i, (task, progress) in enumerate(steps):
                logging.info(f"Loading: {task}")
                if i == 0:
                    self.conn = sqlite3.connect(DB_FILE)
                    self.c = self.conn.cursor()
                    self.c.execute('''CREATE TABLE IF NOT EXISTS transkrip (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        pembicara TEXT,
                        teks TEXT,
                        waktu TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )''')
                    self.conn.commit()
                    init_speaker_db()
                elif i == 1:
                    self.meeting_analyzer = MeetingAnalyzer()
                elif i == 2:
                    self.full_meeting_analyzer = FullMeetingAnalyzer()
                self.progress_var.set(progress)
                time.sleep(0.1)  # Simulate async update
            self.load_complete = True
            logging.info("All resources loaded successfully")
        except Exception as e:
            logging.error(f"Failed to load resources: {str(e)}", exc_info=True)
            self.load_complete = True
            self.window.after(0, lambda: messagebox.showerror("Error", f"Gagal memuat aplikasi: {str(e)}"))

    def update_progress(self):
        if self.load_complete:
            self.window.destroy()
            self.on_complete(self.conn, self.meeting_analyzer, self.full_meeting_analyzer)
        else:
            self.window.after(100, self.update_progress)

# Main Application
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.db_queue = queue.Queue()
        self.db_thread = None
        self.withdraw()  # Hide main window during splash
        import ctypes
        myappid = 'artainovasipersada.ayanotulen.1.0'
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
        os.environ['PYTHONWINDOWICON'] = os.path.abspath('AYA.ico')
        SplashScreen(self, self.initialize_app)

        self.setup_database()

    def setup_database(self):
        """Initialize database components"""
        # Start database worker thread
        self.db_thread = threading.Thread(target=self._db_worker, daemon=True)
        self.db_thread.start()

        # Initialize database tables through the queue
        self.db_queue.put(("init", None))

    def _db_worker(self):
        """Database worker thread that handles all SQLite operations"""
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()

        while True:
            try:
                task = self.db_queue.get()
                if task is None:  # Exit signal
                    break

                func, args = task

                if func == "init":
                    # Initialize tables
                    c.execute('''CREATE TABLE IF NOT EXISTS transkrip (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        pembicara TEXT,
                        teks TEXT,
                        waktu TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )''')
                    conn.commit()

                elif func == "insert":
                    pembicara, teks = args
                    c.execute("INSERT INTO transkrip (pembicara, teks) VALUES (?,?)",
                              (pembicara, teks))
                    conn.commit()

                elif func == "merge":
                    main, sec = args
                    c.execute("UPDATE transkrip SET pembicara=? WHERE pembicara=?",
                              (main, sec))
                    conn.commit()

            except Exception as e:
                print(f"Database error: {e}")
                conn.rollback()
            finally:
                self.db_queue.task_done()

        conn.close()
    def initialize_app(self, conn, meeting_analyzer, full_meeting_analyzer):
        self.conn = conn
        self.meeting_analyzer = meeting_analyzer
        self.full_meeting_analyzer = full_meeting_analyzer
        self.transcript_changed = False
        self.title("AYA Notulen Rapat")
        self.geometry("1050x700")
        try:
            self.iconbitmap(self._get_resource_path('AYA.ico'))
        except Exception as e:
            logging.warning(f"Failed to load app icon: {str(e)}")

        self.setup_menu()
        top = ttk.Frame(self)
        top.pack(fill=tk.X, padx=10, pady=8)
        ttk.Label(top, text="Input:").pack(side=tk.LEFT)
        self.btn_select_device = ttk.Button(top, text="Pilih Device...", command=self.select_device)
        self.btn_select_device.pack(side=tk.LEFT, padx=6)
        self.lbl_device = ttk.Label(top, text="Default Input")
        self.lbl_device.pack(side=tk.LEFT, padx=6)
        self.btn_start = ttk.Button(top, text="Start", command=self.start_recording)
        self.btn_start.pack(side=tk.LEFT, padx=6)
        self.btn_stop = ttk.Button(top, text="Stop", command=self.stop_recording, state="disabled")
        self.btn_stop.pack(side=tk.LEFT, padx=6)
        mid = ttk.Frame(self)
        mid.pack(fill=tk.BOTH, expand=True, padx=10, pady=8)
        left = ttk.Frame(mid, width=250)
        left.pack(side=tk.LEFT, fill=tk.Y)
        left.pack_propagate(False)
        ttk.Label(left, text="Pembicara (klik untuk rename):").pack(anchor="w")
        self.list_speakers = tk.Listbox(left, activestyle='none')
        self.list_speakers.pack(fill=tk.BOTH, expand=True, pady=4)
        self.list_speakers.bind("<Double-1>", self.rename_selected_speaker)
        self.btn_rename = ttk.Button(left, text="Rename Selected", command=self.rename_selected_speaker)
        self.btn_rename.pack(pady=4)
        right = ttk.Frame(mid)
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10)
        ttk.Label(right, text="Notulen (realtime):").pack(anchor="w")
        self.txt = tk.Text(right, wrap=tk.WORD)
        self.txt.pack(fill=tk.BOTH, expand=True, pady=4)
        bottom = ttk.Frame(self)
        bottom.pack(fill=tk.X, padx=10, pady=8)
        ttk.Label(bottom, text="Judul Rapat:").pack(side=tk.LEFT)
        self.ent_judul = ttk.Entry(bottom, width=40)
        self.ent_judul.pack(side=tk.LEFT, padx=6)
        ttk.Label(bottom, text="Peserta:").pack(side=tk.LEFT)
        self.ent_peserta = ttk.Entry(bottom, width=40)
        self.ent_peserta.pack(side=tk.LEFT, padx=6)
        self.btn_export_pdf = ttk.Button(bottom, text="Export PDF", command=self.run_export_pdf)
        self.btn_export_pdf.pack(side=tk.RIGHT, padx=6)
        self.btn_merge = ttk.Button(bottom, text="Gabungkan Speaker", command=self.merge_speakers_gui)
        self.btn_merge.pack(side=tk.RIGHT, padx=6)
        self.btn_export_word = ttk.Button(bottom, text="Export Word", command=self.run_export_word)
        self.btn_export_word.pack(side=tk.RIGHT, padx=6)
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(bottom, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=6)
        self.engine = None
        self.current_device_idx, self.current_device_name = AudioDeviceManager.get_default_input()
        self.lbl_device.config(text=self.current_device_name)
        self.exporter = MeetingExporter(
            engine=self.engine,
            ent_judul=self.ent_judul,
            ent_peserta=self.ent_peserta,
            render_full_transcript_text=self.render_full_transcript_text,
            progress_callback=self.update_progress_bar
        )
        self.deiconify()  # Show main window

    def update_progress_bar(self, percentage):
        self.progress_var.set(percentage)
        self.update_idletasks()

    def _get_resource_path(self, filename):
        base_path = os.path.dirname(os.path.abspath(__file__))
        possible_paths = [
            os.path.join(base_path, filename),
            os.path.join(getattr(sys, '_MEIPASS', ''), filename),
            filename
        ]
        for path in possible_paths:
            if path and os.path.exists(path):
                return path
        raise FileNotFoundError(f"File {filename} not found")

    def setup_menu(self):
        menubar = tk.Menu(self)
        device_menu = tk.Menu(menubar, tearoff=0)
        device_menu.add_command(label="Pilih Input Device...", command=self.select_device)
        device_menu.add_separator()
        input_devices = AudioDeviceManager.get_input_devices()
        for idx, name in input_devices:
            device_menu.add_command(
                label=f"🎤 {name[:30]}...",
                command=lambda i=idx, n=name: self.set_device(i, n)
            )
        loopback_devices = AudioDeviceManager.get_loopback_devices()
        if loopback_devices:
            device_menu.add_separator()
            for idx, name in loopback_devices:
                device_menu.add_command(
                    label=f"🔁 {name[:30]}...",
                    command=lambda i=idx, n=name: self.set_device(i, n)
                )
        menubar.add_cascade(label="Device", menu=device_menu)
        self.config(menu=menubar)

    def select_device(self):
        devices = AudioDeviceManager.get_input_devices() + AudioDeviceManager.get_loopback_devices()
        if not devices:
            messagebox.showerror("Error", "Tidak ada audio input device yang tersedia")
            return
        dialog = tk.Toplevel(self)
        dialog.title("Pilih Audio Input Device")
        dialog.geometry("600x400")
        frame = ttk.Frame(dialog)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        ttk.Label(frame, text="Pilih Device Input:").pack(anchor="w")
        lb = tk.Listbox(frame)
        lb.pack(fill=tk.BOTH, expand=True, pady=5)
        for idx, name in devices:
            prefix = "🔁 " if "(loopback)" in name.lower() else "🎤 "
            lb.insert(tk.END, f"{prefix}{name}")
        def on_select():
            selection = lb.curselection()
            if selection:
                idx = selection[0]
                device_idx, device_name = devices[idx]
                self.set_device(device_idx, device_name)
                dialog.destroy()
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Button(btn_frame, text="Pilih", command=on_select).pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="Batal", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)

    def set_device(self, device_idx, device_name):
        self.current_device_idx = device_idx
        self.current_device_name = device_name
        self.lbl_device.config(text=device_name)
        messagebox.showinfo("Device Dipilih", f"Input device: {device_name}")

    def start_recording(self):
        if self.engine and self.engine.meeting.running:
            logging.warning("Recording already running")
            return
        try:
            logging.info(f"Starting recording with device {self.current_device_idx}: {self.current_device_name}")
            self.engine = RealTimeDiarizationTranscriber(self.current_device_idx, self)
            self.exporter.engine = self.engine
            self.engine.start()
            self.btn_start.config(state="disabled")
            self.btn_stop.config(state="normal")
            self.append_log(f"--- Rekaman dimulai ({self.current_device_name}) ---\n")
            self._refresh_id = self.after(500, self.refresh_transcript, self.engine.meeting)
        except Exception as e:
            logging.error(f"Failed to start recording: {str(e)}", exc_info=True)
            messagebox.showerror("Error", f"Gagal memulai rekaman:\n{str(e)}")

    def stop_recording(self):
        if self.engine:
            try:
                self.engine.stop()
                if hasattr(self, '_refresh_id'):
                    self.after_cancel(self._refresh_id)

                # Add utterances to queue
                for utt in self.engine.meeting.utterances:
                    pemb = self.engine.meeting.speaker_map.get(utt.speaker_id, utt.speaker_id)
                    self.db_queue.put(("insert", (pemb, utt.text)))

                self.btn_start.config(state="normal")
                self.btn_stop.config(state="disabled")
                self.append_log("\n--- Rekaman dihentikan ---\n")

            except Exception as e:
                messagebox.showerror("Error", f"Gagal menghentikan: {str(e)}")
            finally:
                if hasattr(self.engine, 'meeting') and self.engine.meeting.running:
                    self.engine.meeting.running = False

    def append_log(self, text):
        self.txt.insert(tk.END, text)
        self.txt.see(tk.END)

    def render_full_transcript_text(self, meeting):
        lines = []
        for utt in meeting.utterances:
            name = meeting.speaker_map.get(utt.speaker_id, utt.speaker_id)
            lines.append(f"{name}: {utt.text}")
        return "\n".join(lines)

    def refresh_transcript(self, meeting):
        if not meeting.running:
            logging.debug("Meeting stopped, canceling refresh")
            return
        if not self.transcript_changed:
            self._refresh_id = self.after(500, self.refresh_transcript, meeting)
            return
        full = self.render_full_transcript_text(meeting)
        self.txt.delete("1.0", tk.END)
        self.txt.insert(tk.END, full)
        self.txt.see(tk.END)
        self.update_speaker_list(meeting.speaker_map)
        self.transcript_changed = False
        self._refresh_id = self.after(500, self.refresh_transcript, meeting)

    def update_speaker_list(self, speaker_map):
        self.list_speakers.delete(0, tk.END)
        sorted_map = sorted(speaker_map.items(), key=lambda x: int(x[0].replace('-varian', '').split()[-1]))
        for sid, name in sorted_map:
            self.list_speakers.insert(tk.END, f"{sid} → {name}")
        logging.debug(f"Updated speakers: {speaker_map}")

    def rename_selected_speaker(self, event=None):
        if not self.engine:
            return
        idx = self.list_speakers.curselection()
        if not idx:
            return
        line = self.list_speakers.get(idx[0])
        sid = line.split("→")[0].strip()
        current = self.engine.meeting.speaker_map.get(sid, sid)
        new_name = simpledialog.askstring("Ganti Nama", f"{sid} jadi:", initialvalue=current, parent=self)
        if new_name and new_name.strip():
            self.engine.meeting.speaker_map[sid] = new_name.strip()
            self.refresh_transcript(self.engine.meeting)

    def run_export_pdf(self):
        def export_task():
            try:
                # Create new event loop for this thread
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)

                # Ensure exporter has current engine reference
                self.exporter.engine = self.engine

                # Run the export
                loop.run_until_complete(self.exporter.export_pdf())
                loop.close()
            except Exception as e:
                print(f"Export error: {e}")
                self.after(0, lambda: messagebox.showerror("Export Error", str(e)))

        # Run in separate thread
        threading.Thread(target=export_task, daemon=True).start()

    def run_export_word(self):
        def export_task():
            try:
                # Create new event loop for this thread
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)

                # Ensure exporter has current engine reference
                self.exporter.engine = self.engine

                # Run the export
                loop.run_until_complete(self.exporter.export_word())
                loop.close()
            except Exception as e:
                print(f"Export error: {e}")
                self.after(0, lambda: messagebox.showerror("Export Error", str(e)))

        # Run in separate thread
        threading.Thread(target=export_task, daemon=True).start()

    def _set_app_icon(self, window=None):
        target = window or self
        try:
            target.iconbitmap(self._get_resource_path('AYA.ico'))
        except Exception as e:
            logging.warning(f"Failed to set icon: {e}")

    def merge_speakers_gui(self):
        if not self.engine:
            messagebox.showerror("Error", "Tidak ada data meeting aktif")
            return
        dialog = tk.Toplevel(self)
        dialog.title("Gabungkan Pembicara")
        dialog.geometry("500x350")
        self._set_app_icon(dialog)
        dialog.grab_set()
        main_frame = ttk.Frame(dialog)
        main_frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        ttk.Label(main_frame, text="Speaker Utama:").pack(anchor=tk.W)
        main_speaker_var = tk.StringVar()
        main_cb = ttk.Combobox(
            main_frame,
            textvariable=main_speaker_var,
            values=list(self.engine.meeting.speaker_map.keys()),
            state="readonly"
        )
        main_cb.pack(fill=tk.X, pady=5)
        ttk.Label(main_frame, text="Speaker yang akan digabung:").pack(anchor=tk.W)
        sec_speaker_var = tk.StringVar()
        sec_cb = ttk.Combobox(
            main_frame,
            textvariable=sec_speaker_var,
            values=list(self.engine.meeting.speaker_map.keys()),
            state="readonly"
        )
        sec_cb.pack(fill=tk.X, pady=5)
        info_label = ttk.Label(main_frame, text="", foreground="red")
        info_label.pack(pady=10)
        def update_info(*args):
            main = main_speaker_var.get()
            sec = sec_speaker_var.get()
            if main and sec:
                if main == sec:
                    info_label.config(text="Pilih speaker yang berbeda!", foreground="red")
                else:
                    count = sum(1 for u in self.engine.meeting.utterances if u.speaker_id == sec)
                    info_label.config(
                        text=f"Semua {count} ucapan dari '{sec}' akan digabung ke '{main}'",
                        foreground="green"
                    )
            else:
                info_label.config(text="Pilih kedua speaker", foreground="red")
        main_speaker_var.trace_add('write', update_info)
        sec_speaker_var.trace_add('write', update_info)
        def do_merge():
            main = main_speaker_var.get()
            sec = sec_speaker_var.get()
            if not main or not sec:
                messagebox.showerror("Error", "Pilih kedua speaker")
                return
            if main == sec:
                messagebox.showerror("Error", "Pilih speaker yang berbeda")
                return
            try:
                merge_speakers(main, sec)
                for utterance in self.engine.meeting.utterances:
                    if utterance.speaker_id == sec:
                        utterance.speaker_id = main
                if sec in self.engine.meeting.speaker_map:
                    del self.engine.meeting.speaker_map[sec]
                self.refresh_transcript(self.engine.meeting)
                messagebox.showinfo("Sukses", f"Semua ucapan dari '{sec}' telah digabung ke '{main}'")
                dialog.destroy()
            except Exception as e:
                messagebox.showerror("Error", f"Gagal menggabungkan: {str(e)}")
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=10)
        ttk.Button(button_frame, text="Batal", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Gabungkan", command=do_merge, style='Accent.TButton').pack(side=tk.RIGHT, padx=5)

    def auto_merge_similar_speakers(self, threshold=0.8):
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        try:
            c.execute("SELECT name, embedding FROM speaker_profiles ORDER BY usage_count DESC")
            speakers = c.fetchall()
            merged = set()
            for i, (name1, emb1) in enumerate(speakers):
                if name1 in merged:
                    continue
                emb1 = np.frombuffer(emb1, dtype=np.float32)
                emb1 = emb1 / np.linalg.norm(emb1)
                for j, (name2, emb2) in enumerate(speakers[i + 1:], i + 1):
                    if name2 in merged:
                        continue
                    emb2 = np.frombuffer(emb2, dtype=np.float32)
                    emb2 = emb2 / np.linalg.norm(emb2)
                    sim = np.dot(emb1, emb2)
                    if sim >= threshold:
                        c.execute("SELECT usage_count FROM speaker_profiles WHERE name=?", (name1,))
                        cnt1 = c.fetchone()[0]
                        c.execute("SELECT usage_count FROM speaker_profiles WHERE name=?", (name2,))
                        cnt2 = c.fetchone()[0]
                        main, sec = (name1, name2) if cnt1 >= cnt2 else (name2, name1)
                        merge_speakers(main, sec)
                        merged.add(sec)
                        break
            conn.commit()
            return True
        except Exception as e:
            conn.rollback()
            raise e
        finally:
            conn.close()

    def show_merge_suggestions(self, threshold=0.75):
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        c.execute("""
            SELECT name, embedding, 
                   (SELECT COUNT(*) FROM transkrip WHERE pembicara=name) as count
            FROM speaker_profiles
            ORDER BY count DESC
        """)
        speakers = c.fetchall()
        similarities = []
        embeddings = []
        for name, emb_bytes, _ in speakers:
            emb = np.frombuffer(emb_bytes, dtype=np.float32)
            emb = emb / np.linalg.norm(emb)
            embeddings.append((name, emb))
        for i, (name1, emb1) in enumerate(embeddings):
            for j, (name2, emb2) in enumerate(embeddings[i + 1:], i + 1):
                sim = np.dot(emb1, emb2)
                if sim >= threshold:
                    similarities.append((name1, name2, sim))
        similarities.sort(key=lambda x: -x[2])
        dialog = tk.Toplevel(self)
        dialog.title("Saran Penggabungan Pembicara")
        dialog.geometry("700x500")
        self._set_app_icon(dialog)
        frame = ttk.Frame(dialog)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        ttk.Label(frame, text="Pasangan pembicara yang mungkin sama:").pack(anchor="w")
        tree = ttk.Treeview(frame, columns=("speaker1", "speaker2", "similarity"), show="headings")
        tree.heading("speaker1", text="Pembicara 1")
        tree.heading("speaker2", text="Pembicara 2")
        tree.heading("similarity", text="Kemiripan")
        tree.column("speaker1", width=200)
        tree.column("speaker2", width=200)
        tree.column("similarity", width=100)
        for name1, name2, sim in similarities:
            tree.insert("", tk.END, values=(name1, name2, f"{sim:.3f}"))
        tree.pack(fill=tk.BOTH, expand=True, pady=5)
        def on_merge():
            selected = tree.focus()
            if not selected:
                return
            item = tree.item(selected)
            name1, name2, _ = item['values']
            c.execute("SELECT COUNT(*) FROM transkrip WHERE pembicara=?", (name1,))
            cnt1 = c.fetchone()[0]
            c.execute("SELECT COUNT(*) FROM transkrip WHERE pembicara=?", (name2,))
            cnt2 = c.fetchone()[0]
            main, sec = (name1, name2) if cnt1 >= cnt2 else (name2, name1)
            if messagebox.askyesno(
                    "Konfirmasi",
                    f"Gabungkan semua ucapan {sec} ke {main}?\n"
                    f"Tindakan ini tidak dapat dibatalkan."
            ):
                try:
                    merge_speakers(main, sec)
                    if sec in self.engine.meeting.speaker_map:
                        del self.engine.meeting.speaker_map[sec]
                    self.refresh_transcript(self.engine.meeting)
                    tree.delete(selected)
                    messagebox.showinfo("Sukses", f"Pembicara {sec} telah digabung ke {main}")
                except Exception as e:
                    messagebox.showerror("Error", str(e))
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Button(btn_frame, text="Gabungkan", command=on_merge).pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="Tutup", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)

# RealTimeDiarizationTranscriber (unchanged except for logging)
class RealTimeDiarizationTranscriber:
    def __init__(self, device_index: Optional[int], gui_ref):
        logging.info(f"Initializing RealTimeDiarizationTranscriber with device_index={device_index}")
        self.max_queue_size = 100
        self.force_stop = False
        self.audio_lock = threading.Lock()
        self.gui = gui_ref
        self.device_index = device_index
        self.q = queue.Queue()
        self.vad = webrtcvad.Vad(VAD_MODE)
        self.sd_stream = None
        self.worker_thread = None
        self.stop_event = threading.Event()
        self.meeting = MeetingState()
        try:
            spkr_model_path = self._get_valid_model_path('ecapa_voxceleb', ['hyperparams.yaml', 'embedding_model.ckpt'])
            self.spkr_model = EncoderClassifier.from_hparams(source=spkr_model_path, run_opts={"device": WHISPER_DEVICE})
            logging.info(f"Speaker model loaded: {spkr_model_path}")
            whisper_path = self._get_valid_model_path('whisper_medium', ['model.bin', 'config.json'])
            self.whisper = WhisperModel(whisper_path, device=WHISPER_DEVICE, compute_type=WHISPER_COMPUTE_TYPE)
            logging.info(f"Whisper model loaded: {whisper_path}")
        except Exception as e:
            logging.error(f"Failed to initialize models: {str(e)}", exc_info=True)
            raise
        self.frame_bytes = int(SAMPLE_RATE * FRAME_MS / 1000)
        self.curr_frames = []
        self.non_speech_run = 0
        self.audio_clock = 0.0
        self.min_audio_energy = 0.01
        self.last_valid_utterance_time = 0

    def _get_valid_model_path(self, model_name, required_files):
        possible_paths = [
            os.path.join(os.path.dirname(__file__), 'models', model_name),
            os.path.join(getattr(sys, '_MEIPASS', ''), 'models', model_name),
            os.path.join('models', model_name)
        ]
        for path in possible_paths:
            if path and os.path.exists(path):
                if all(os.path.exists(os.path.join(path, f)) for f in required_files):
                    return path
                logging.warning(f"Model {model_name} missing files at {path}")
        raise FileNotFoundError(f"Model {model_name} not found")

    def start(self):
        self.meeting = MeetingState(running=True, start_time=datetime.datetime.now())
        self.stop_event.clear()
        self.sd_stream = sd.InputStream(
            device=self.device_index,
            channels=1,
            samplerate=SAMPLE_RATE,
            blocksize=self.frame_bytes,
            dtype='int16',
            callback=self._callback
        )
        self.sd_stream.start()
        self.worker_thread = threading.Thread(target=self._worker_loop, daemon=True)
        self.worker_thread.start()

    def stop(self):
        logging.info("Initiating shutdown")
        self.force_stop = True
        self.stop_event.set()
        if self.sd_stream:
            try:
                with self.audio_lock:
                    self.sd_stream.stop()
                    self.sd_stream.close()
                    self.sd_stream = None
            except Exception as e:
                logging.warning(f"Error closing stream: {e}")
        try:
            while not self.q.empty():
                self.q.get_nowait()
        except Exception as e:
            logging.warning(f"Error clearing queue: {e}")
        if self.worker_thread and self.worker_thread.is_alive():
            self.worker_thread.join(timeout=2.0)
            if self.worker_thread.is_alive():
                logging.warning("Worker thread not responding")
        self.meeting.running = False
        self.meeting.end_time = datetime.datetime.now()
        logging.info("Shutdown complete")

    def _callback(self, indata, frames, time_info, status):
        if status:
            logging.warning(f"Audio status: {status}")
        try:
            self.q.put(indata.copy())
        except Exception as e:
            logging.error(f"Error in callback: {str(e)}", exc_info=True)

    def _extract_embedding(self, wav_float_16k: torch.Tensor) -> torch.Tensor:
        with torch.no_grad():
            emb = self.spkr_model.encode_batch(wav_float_16k).squeeze(0).squeeze(0)
            emb = F.normalize(emb, p=2, dim=0)
        return emb

    def _assign_speaker(self, emb: torch.Tensor) -> str:
        emb = F.normalize(emb, p=2, dim=0)
        best_sim = 0.0
        best_speaker = None
        threshold = self.get_dynamic_threshold()
        for spk_id, spk_emb in self.meeting.speaker_embeds.items():
            sim = torch.dot(emb, spk_emb).item()
            logging.debug(f"Similarity to {spk_id}: {sim:.3f} (threshold: {threshold:.3f})")
            if sim > best_sim:
                best_sim = sim
                best_speaker = spk_id
        if best_speaker and best_sim >= threshold:
            updated_emb = 0.9 * self.meeting.speaker_embeds[best_speaker] + 0.1 * emb
            updated_emb = F.normalize(updated_emb, p=2, dim=0)
            self.meeting.speaker_embeds[best_speaker] = updated_emb
            return best_speaker
        else:
            if best_sim > 0.60:
                new_name = f"{best_speaker}-varian"
            else:
                new_name = f"Pembicara {self.meeting.next_speaker_idx}"
                self.meeting.next_speaker_idx += 1
            self.meeting.speaker_embeds[new_name] = emb
            if new_name not in self.meeting.speaker_map:
                self.meeting.speaker_map[new_name] = new_name
            logging.debug(f"Assigned speaker: {new_name}")
            return new_name

    def get_dynamic_threshold(self):
        if len(self.meeting.speaker_embeds) < 3:
            return 0.65
        similarities = []
        emb_list = list(self.meeting.speaker_embeds.values())
        for i in range(len(emb_list)):
            for j in range(i + 1, len(emb_list)):
                sim = torch.dot(emb_list[i], emb_list[j]).item()
                similarities.append(sim)
        threshold = np.percentile(similarities, 75) + 0.1
        return min(max(threshold, 0.60), 0.85)

    def _is_valid_audio(self, audio_data):
        energy = np.mean(np.abs(audio_data))
        logging.debug(f"Checking audio validity: energy={energy}, threshold={self.min_audio_energy}")
        return energy > self.min_audio_energy

    def _transcribe_float32(self, floataudio: np.ndarray) -> str:
        if not self._is_valid_audio(floataudio):
            logging.debug("Transcription skipped: invalid audio")
            return ""
        logging.info("Starting transcription")
        torch.cuda.empty_cache()
        try:
            segments, _ = self.whisper.transcribe(
                floataudio,
                beam_size=5,
                vad_filter=True,
                language="id",
                no_speech_threshold=0.3,
                compression_ratio_threshold=1.8
            )
            text = " ".join(seg.text.strip() for seg in segments).strip()
            logging.info(f"Transcription completed: {text}")
            torch.cuda.empty_cache()
            return text
        except Exception as e:
            logging.error(f"Transcription error: {str(e)}", exc_info=True)
            torch.cuda.empty_cache()
            return ""

    def _finalize_current_segment(self):
        if not self.curr_frames:
            logging.debug("No frames to finalize")
            return
        seg_int16 = np.concatenate(self.curr_frames, axis=0)
        dur = len(seg_int16) / SAMPLE_RATE
        logging.debug(f"Segment duration: {dur:.2f}s")
        if not self._is_valid_audio(seg_int16.astype(np.float32) / 32768.0):
            logging.debug("Segment invalid: low energy")
            self.curr_frames = []
            return
        if dur < MIN_UTTER_SEC:
            logging.debug(f"Segment too short: {dur:.2f}s < {MIN_UTTER_SEC}s")
            self.curr_frames = []
            return
        if dur > MAX_UTTER_SEC:
            logging.debug(f"Segment too long: {dur:.2f}s, truncating to {MAX_UTTER_SEC}s")
            seg_int16 = seg_int16[:int(MAX_UTTER_SEC * SAMPLE_RATE)]
        seg_float32 = seg_int16.astype(np.float32) / 32768.0
        wav_tensor = torch.from_numpy(seg_float32).unsqueeze(0).to(WHISPER_DEVICE)
        text = self._transcribe_float32(seg_float32)
        if not text.strip():
            logging.debug("No transcription produced")
            self.curr_frames = []
            return
        emb = self._extract_embedding(wav_tensor)
        spk_id = self._assign_speaker(emb)
        end_time = self.audio_clock
        start_time = max(0.0, end_time - len(seg_int16) / SAMPLE_RATE)
        if self.meeting.utterances:
            last_utt = self.meeting.utterances[-1]
            if last_utt.speaker_id == spk_id and (start_time - last_utt.end_time) < 2.0:
                last_utt.text += " " + text
                last_utt.end_time = end_time
                logging.debug(f"Merged utterance for {spk_id}")
            else:
                utt = Utterance(spk_id, text, start_time, end_time)
                self.meeting.utterances.append(utt)
                logging.debug(f"New utterance for {spk_id}")
        else:
            utt = Utterance(spk_id, text, start_time, end_time)
            self.meeting.utterances.append(utt)
            logging.debug(f"First utterance for {spk_id}")
        self.gui.transcript_changed = True
        self.curr_frames = []

    def _worker_loop(self):
        max_empty_cycles = 20
        empty_cycles = 0
        logging.info("Starting worker loop")
        while not self.stop_event.is_set() and empty_cycles < max_empty_cycles:
            try:
                if self.q.qsize() > self.max_queue_size:
                    logging.warning(f"Queue full ({self.q.qsize()}/{self.max_queue_size}), dropping oldest frame")
                    self.q.get_nowait()
                frames = []
                try:
                    while len(frames) < 5:
                        frame = self.q.get(timeout=0.05)
                        frames.append(frame)
                except queue.Empty:
                    pass
                empty_cycles = 0 if frames else empty_cycles + 1
                if frames:
                    with self.audio_lock:
                        for frame in frames:
                            self.audio_clock += FRAME_MS / 1000.0
                            is_speech = self.vad.is_speech(frame.tobytes(), SAMPLE_RATE)
                            logging.debug(f"Frame processed, is_speech={is_speech}, queue_size={self.q.qsize()}")
                            if is_speech:
                                self.curr_frames.append(frame.squeeze())
                                self.non_speech_run = 0
                            elif self.curr_frames:
                                self.non_speech_run += 1
                                if self.non_speech_run >= SILENCE_TAIL_FRAMES:
                                    logging.info("Finalizing segment")
                                    self._finalize_current_segment()
            except Exception as e:
                logging.error(f"Worker loop error: {str(e)}", exc_info=True)
                break
        logging.info("Worker thread exiting")
        with self.audio_lock:
            if self.curr_frames and not self.stop_event.is_set():
                logging.info("Final cleanup: finalizing remaining segment")
                self._finalize_current_segment()

def main():
    init_speaker_db()
    app = App()
    app.mainloop()

if __name__ == "__main__":
    main()